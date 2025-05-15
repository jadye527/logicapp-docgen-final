
from docx import Document
from docx.shared import Inches
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text=None):
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def generate_document(
    architecture, execution, flow_text, data_flow, hybrid_text, conditions,
    overview_summary=None,
    architecture_summary=None,
    execution_summary=None,
    purpose_summary=None,
    security_summary=None,
    error_handling_summary=None,
    integration_endpoints=None,
    docx_template=None
):
    doc = Document(docx_template)

    logic_app_name = architecture.get("name", "LogicApp")
    region = architecture.get("location", "unknown")
    resource_group = architecture.get("resource_group", "n/a")
    subscription_id = architecture.get("subscription_id", "n/a")
    automation_account = architecture.get("automation_account", "n/a")
    tags = architecture.get("tags", {})

    doc.add_heading(f"Azure Logic App Workflow – {logic_app_name}", level=1)

    doc.add_heading("Overview", level=2)
    if overview_summary:
        doc.add_paragraph(overview_summary)

    doc.add_heading("Purpose and Function", level=2)
    if purpose_summary:
        doc.add_paragraph(purpose_summary)
    doc.add_heading("Key Responsibilities", level=3)
    for item in execution:
        doc.add_paragraph(item, style='Bullets')

    doc.add_heading("Architecture", level=2)
    doc.add_heading("Overview", level=3)
    if architecture_summary:
        doc.add_paragraph(architecture_summary)

    doc.add_heading("Deployment Metadata", level=3)
    doc.add_paragraph(f"Logic App Name: {logic_app_name}", style='Bullets')
    doc.add_paragraph(f"Resource Group: {resource_group}", style='Bullets')
    doc.add_paragraph(f"Subscription ID: {subscription_id}", style='Bullets')
    doc.add_paragraph(f"Location: {region}", style='Bullets')
    doc.add_paragraph(f"Automation Account: {automation_account}", style='Bullets')
    for k, v in tags.items():
        doc.add_paragraph(f"Tag - {k}: {v}", style='Bullets')

    doc.add_heading("Execution Logic Summary", level=2)
    if execution_summary:
        doc.add_paragraph(execution_summary)

    doc.add_heading("Logic App Flow Diagram", level=2)
    flow_png = os.path.join("output", f"{logic_app_name}_Flow.png")
    if os.path.exists(flow_png):
        doc.add_picture(flow_png, width=Inches(6.0))
        doc.paragraphs[-1].alignment = 1
        caption = doc.add_paragraph(style="Caption")
        caption.add_run("Figure 1 – Azure Logic App Flow Diagram")

    else:
        doc.add_paragraph("❌ Could not render Logic App Flow Diagram.")

    doc.add_heading("Data Flow Diagram", level=2)
    for item in data_flow:
        doc.add_paragraph(item, style='Bullets')

    doc.add_heading("Hybrid Integration Diagram", level=2)
    hybrid_png = os.path.join("output", f"{logic_app_name}_Hybrid.png")
    if os.path.exists(hybrid_png):
        doc.add_picture(hybrid_png, width=Inches(6.0))
        doc.paragraphs[-1].alignment = 1
        caption = doc.add_paragraph(style="Caption")
        caption.add_run("Figure 2 – Hybrid Integration Diagram")

    else:
        doc.add_paragraph("❌ Could not render Hybrid Integration Diagram.")
    for item in hybrid_text:
        doc.add_paragraph(item, style='Bullets')

    doc.add_heading("Security and Authentication", level=2)
    if security_summary:
        doc.add_paragraph(security_summary)

    doc.add_heading("Error Handling", level=2)
    if error_handling_summary:
        doc.add_paragraph(error_handling_summary)

    doc.add_heading("Appendix: Integration Endpoints", level=2)
    if integration_endpoints:
        for item in integration_endpoints:
            if "http" in item:
                label, url = item.split(":", 1)
                p = doc.add_paragraph(style='Bullets')
                add_hyperlink(p, url.strip(), f"{label.strip()}: {url.strip()}")
            else:
                doc.add_paragraph(item, style='Bullets')

    return doc
